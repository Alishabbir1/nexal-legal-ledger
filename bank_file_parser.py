"""
Bank statement file parser - CSV, XLSX, PDF, DOCX.
Extracts date, description, amount, reference for reconciliation matching.
PDF/DOCX: attempts structured extraction; falls back to Manual Review Mode.
"""
import csv
import io
import re
from datetime import datetime
from typing import List, Dict, Optional, Tuple
from decimal import Decimal, InvalidOperation

# Optional imports for XLSX, PDF, DOCX
try:
    import openpyxl
except ImportError:
    openpyxl = None
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'pdf', 'docx'}
STRUCTURED_FORMATS = {'csv', 'xlsx'}

# Column name variants for flexible matching
DATE_COLUMNS = {'date', 'transaction_date', 'value_date', 'posted_date', 'transaction date'}
AMOUNT_COLUMNS = {'amount', 'debit', 'credit', 'value', 'sum', 'total'}
DESCRIPTION_COLUMNS = {'description', 'narrative', 'details', 'particulars', 'memo'}
REFERENCE_COLUMNS = {'reference', 'ref', 'transaction_ref', 'cheque_number', 'cheque_no'}


def _normalize_column(name: str) -> str:
    return (name or '').strip().lower().replace(' ', '_').replace('-', '_')


def _find_column(row_or_headers: dict, candidates: set) -> Optional[str]:
    """Find first matching column from dict keys."""
    keys = {_normalize_column(k): k for k in row_or_headers}
    for c in candidates:
        if c in keys:
            return keys[c]
    return None


def _parse_amount(val) -> Optional[Decimal]:
    if val is None or val == '':
        return None
    s = str(val).strip()
    s = re.sub(r'[^\d.\-]', '', s)
    if not s:
        return None
    try:
        return Decimal(s)
    except (InvalidOperation, ValueError):
        return None


def _parse_date(val) -> Optional[str]:
    if val is None or val == '':
        return None
    if hasattr(val, 'strftime'):
        return val.strftime('%Y-%m-%d')
    s = str(val).strip()
    for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y', '%d.%m.%Y', '%Y/%m/%d', '%d/%m/%y'):
        try:
            dt = datetime.strptime(s[:10], fmt)
            return dt.strftime('%Y-%m-%d')
        except ValueError:
            continue
    return None


def _extract_ref_from_description(desc: str) -> str:
    """Try to extract reference-like pattern from description."""
    if not desc:
        return ''
    # Cheque number pattern
    m = re.search(r'\b(\d{4,})\b', desc)
    if m:
        return m.group(1)
    return ''


def parse_csv(content: bytes, filename: str = '') -> Tuple[List[Dict], bool, Optional[str]]:
    """
    Parse CSV. Returns (entries, manual_review, error).
    manual_review True if we couldn't extract structured data.
    """
    try:
        text = content.decode('utf-8-sig', errors='replace')
        reader = csv.DictReader(io.StringIO(text))
        rows = list(reader)
    except Exception as e:
        return [], False, f"Could not read CSV: {e}"
    if not rows:
        return [], True, "File is empty"
    headers = {_normalize_column(h): h for h in rows[0].keys()}
    date_col = _find_column(headers, DATE_COLUMNS)
    amt_col = _find_column(headers, AMOUNT_COLUMNS)
    if not date_col or not amt_col:
        return [], True, "Required columns (date, amount) not found. Use date, description, amount, reference."
    desc_col = _find_column(headers, DESCRIPTION_COLUMNS) or list(rows[0].keys())[1] if len(rows[0]) > 1 else None
    ref_col = _find_column(headers, REFERENCE_COLUMNS)
    entries = []
    for i, row in enumerate(rows):
        d = _parse_date(row.get(date_col))
        amt = _parse_amount(row.get(amt_col))
        if d is None and amt is None:
            continue
        desc = (row.get(desc_col) or row.get('description', '') or '').strip()
        ref = (row.get(ref_col) or '').strip() if ref_col else _extract_ref_from_description(desc)
        if amt is not None:
            entries.append({
                'date': d or '',
                'description': desc,
                'amount': float(amt),
                'reference': ref or desc[:50],
                'row': i + 1,
            })
    if not entries:
        return [], True, "No valid transaction rows found"
    return entries, False, None


def parse_xlsx(content: bytes, filename: str = '') -> Tuple[List[Dict], bool, Optional[str]]:
    """Parse XLSX. Returns (entries, manual_review, error)."""
    if openpyxl is None:
        return [], False, "openpyxl not installed. Install with: pip install openpyxl"
    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        wb.close()
    except Exception as e:
        return [], False, f"Could not read XLSX: {e}"
    if not rows:
        return [], True, "Sheet is empty"
    headers = [str(h or '').strip() for h in rows[0]]
    hnorm = {_normalize_column(h): i for i, h in enumerate(headers)}
    def col(name_set):
        for c in name_set:
            if c in hnorm:
                return hnorm[c]
        return None
    date_idx = col(DATE_COLUMNS)
    amt_idx = col(AMOUNT_COLUMNS)
    if date_idx is None or amt_idx is None:
        return [], True, "Required columns (date, amount) not found."
    desc_idx = col(DESCRIPTION_COLUMNS) or (1 if len(headers) > 1 else None)
    ref_idx = col(REFERENCE_COLUMNS)
    entries = []
    for i, row in enumerate(rows[1:], start=2):
        if not row:
            continue
        r = list(row) + [None] * (max((date_idx or 0), (amt_idx or 0), (desc_idx or 0), (ref_idx or 0)) + 1 - len(row))
        d = _parse_date(r[date_idx] if date_idx is not None else None)
        amt = _parse_amount(r[amt_idx] if amt_idx is not None else None)
        if d is None and amt is None:
            continue
        desc = str(r[desc_idx] or '') if desc_idx is not None else ''
        ref = str(r[ref_idx] or '') if ref_idx is not None else _extract_ref_from_description(desc)
        if amt is not None:
            entries.append({
                'date': d or '',
                'description': desc,
                'amount': float(amt),
                'reference': ref or desc[:50],
                'row': i,
            })
    if not entries:
        return [], True, "No valid transaction rows found"
    return entries, False, None


def _extract_structured_from_text(text: str) -> List[Dict]:
    """Attempt to extract date/amount/desc from free text using patterns."""
    entries = []
    lines = text.splitlines()
    # Pattern: date amount description (various formats)
    date_amt = re.compile(
        r'(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4})\s+([\d,]+\.?\d*)\s+(.+)',
        re.IGNORECASE
    )
    for line in lines:
        m = date_amt.search(line)
        if m:
            d = _parse_date(m.group(1))
            amt = _parse_amount(m.group(2))
            desc = m.group(3).strip()[:200]
            if d and amt is not None:
                entries.append({
                    'date': d,
                    'description': desc,
                    'amount': float(amt),
                    'reference': _extract_ref_from_description(desc) or desc[:50],
                    'row': len(entries) + 1,
                })
    return entries


def parse_pdf(content: bytes, filename: str = '') -> Tuple[List[Dict], bool, Optional[str]]:
    """Parse PDF. Extract text and try structured parsing. Fall back to manual review."""
    if PdfReader is None:
        return [], False, "PyPDF2 not installed. Install with: pip install PyPDF2"
    try:
        reader = PdfReader(io.BytesIO(content))
        text = ''
        for page in reader.pages:
            text += page.extract_text() or ''
    except Exception as e:
        return [], False, f"Could not read PDF: {e}"
    if not text.strip():
        return [], True, "No text could be extracted from PDF. File may be scanned/image-based."
    entries = _extract_structured_from_text(text)
    if entries:
        return entries, False, None
    return [], True, "Structured data could not be parsed. File placed in Manual Review Mode."


def parse_docx(content: bytes, filename: str = '') -> Tuple[List[Dict], bool, Optional[str]]:
    """Parse DOCX. Extract text and try structured parsing."""
    if DocxDocument is None:
        return [], False, "python-docx not installed. Install with: pip install python-docx"
    try:
        doc = DocxDocument(io.BytesIO(content))
        text = '\n'.join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += '\n' + ' '.join(cell.text or '' for cell in row.cells)
    except Exception as e:
        return [], False, f"Could not read DOCX: {e}"
    if not text.strip():
        return [], True, "Document appears empty."
    entries = _extract_structured_from_text(text)
    if entries:
        return entries, False, None
    return [], True, "Structured data could not be parsed. File placed in Manual Review Mode."


def parse_bank_file(content: bytes, filename: str) -> Tuple[List[Dict], bool, Optional[str]]:
    """
    Parse bank file by extension. Returns (entries, manual_review, error).
    manual_review=True: display for manual confirmation, do not auto-match.
    """
    ext = (filename.rsplit('.', 1)[-1] or '').lower()
    if ext == 'csv':
        return parse_csv(content, filename)
    if ext == 'xlsx':
        return parse_xlsx(content, filename)
    if ext == 'pdf':
        return parse_pdf(content, filename)
    if ext == 'docx':
        return parse_docx(content, filename)
    return [], False, f"Unsupported format: {ext}. Use .csv, .xlsx, .pdf, or .docx"
